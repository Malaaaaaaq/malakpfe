#!/usr/bin/env python3
"""Generate PFE report for ParLak project in Word (.docx) format."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 5):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x04, 0x15, 0x62)
    h.font.name = 'Calibri'
    if level == 1:
        h.font.size = Pt(16)
    elif level == 2:
        h.font.size = Pt(13)
    elif level == 3:
        h.font.size = Pt(11.5)
    else:
        h.font.size = Pt(11)

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 0.8)
    return p

def add_para(doc, text, bold=False, italic=False, size=11, alignment=None, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if alignment:
        p.alignment = alignment
    return p

# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

add_para(doc, 'Royaume du Maroc', bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x04, 0x15, 0x62))
add_para(doc, 'Office de la Formation Professionnelle et de la Promotion du Travail', bold=False, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x64, 0x74, 0x8b))
add_para(doc, 'ISTA NTIC Hay Salam – Nador', bold=False, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x64, 0x74, 0x8b))

doc.add_paragraph()
doc.add_paragraph()

add_para(doc, 'Projet de Fin d\'Études', bold=True, size=16, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x04, 0x15, 0x62))
add_para(doc, 'Filière : Développement Informatique', bold=False, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

add_para(doc, 'ParLak', bold=True, size=26, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x04, 0x15, 0x62))
add_para(doc, 'Plateforme Intelligente de Stationnement', bold=False, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x3b, 0x82, 0xf6))

doc.add_paragraph()

add_para(doc, 'Réalisé par :', bold=True, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, 'Malak BOUDAD & Fatima BOUGROUN', bold=False, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

add_para(doc, f'Année universitaire : {datetime.date.today().year - 1} / {datetime.date.today().year}', bold=False, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════
doc.add_heading('Table des matières', level=1)
toc = [
    'Introduction générale',
    '',
    'Chapitre 1 : Contexte et présentation du projet',
    '    1.1 Problématique',
    '    1.2 Objectifs du projet',
    '    1.3 Périmètre du projet',
    '',
    'Chapitre 2 : Analyse et spécification des besoins',
    '    2.1 Besoins fonctionnels',
    '    2.2 Besoins non fonctionnels',
    '    2.3 Acteurs du système',
    '    2.4 Diagramme de cas d\'utilisation',
    '',
    'Chapitre 3 : Conception',
    '    3.1 Architecture du système',
    '    3.2 Modélisation de la base de données',
    '    3.3 Diagrammes UML (séquence, classes)',
    '',
    'Chapitre 4 : Implémentation',
    '    4.1 Technologies utilisées',
    '    4.2 Backend (Laravel 12)',
    '    4.3 Frontend (React 19)',
    '',
    'Chapitre 5 : Fonctionnalités et manuel d\'utilisation',
    '    5.1 Présentation des interfaces',
    '    5.2 Manuel d\'utilisation',
    '',
    'Chapitre 6 : Tests et validation',
    '',
    'Conclusion et perspectives',
    '',
    'Bibliographie / Webographie',
    '',
    'Annexes',
]
for item in toc:
    if item == '':
        doc.add_paragraph()
    else:
        p = doc.add_paragraph()
        r = p.add_run(item)
        r.font.size = Pt(11)
        if not item.startswith('    '):
            r.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# INTRODUCTION GENERALE
# ══════════════════════════════════════════════════════════════
doc.add_heading('Introduction générale', level=1)

doc.add_paragraph(
    'La gestion du stationnement urbain est devenue un défi majeur dans les grandes villes marocaines. '
    'Entre la croissance du parc automobile, le manque de places disponibles et la difficulté de trouver '
    'une solution fiable et rapide, les conducteurs perdent un temps considérable chaque jour. '
    'C\'est dans ce contexte que s\'inscrit le projet ParLak.'
)
doc.add_paragraph(
    'ParLak est une plateforme intelligente de stationnement qui vise à simplifier la réservation de places '
    'de parking. Elle met en relation trois acteurs principaux : les clients souhaitant réserver une place, '
    'les agents gérant les parkings au quotidien, et les administrateurs supervisant l\'ensemble du système.'
)
doc.add_paragraph(
    'Ce rapport présente l\'analyse, la conception et la réalisation de cette plateforme. Il est structuré '
    'en six chapitres. Le premier chapitre expose le contexte et les objectifs du projet. Le deuxième '
    'chapitre détaille l\'analyse des besoins. Le troisième chapitre présente la conception architecturale '
    'et la modélisation des données. Le quatrième chapitre décrit l\'implémentation technique. Le cinquième '
    'chapitre présente les fonctionnalités et le manuel d\'utilisation. Enfin, le sixième chapitre aborde '
    'les tests et la validation.'
)

# ══════════════════════════════════════════════════════════════
# CHAPITRE 1
# ══════════════════════════════════════════════════════════════
doc.add_heading('Chapitre 1 : Contexte et présentation du projet', level=1)

doc.add_heading('1.1 Problématique', level=2)
doc.add_paragraph(
    'Dans les grandes villes comme Casablanca, Rabat, Marrakech et Tanger, trouver une place de '
    'stationnement est souvent une source de stress et de perte de temps. Les problèmes identifiés sont :'
)
add_bullet(doc, 'Absence d\'information en temps réel sur les places disponibles')
add_bullet(doc, 'Impossibilité de réserver une place à l\'avance')
add_bullet(doc, 'Gestion manuelle et inefficace des parkings')
add_bullet(doc, 'Manque de communication entre les gestionnaires de parking et les conducteurs')
add_bullet(doc, 'Absence d\'un système de fidélisation et de codes promotionnels')

doc.add_heading('1.2 Objectifs du projet', level=2)
doc.add_paragraph('Les objectifs principaux du projet ParLak sont :')
add_bullet(doc, 'Permettre aux conducteurs de localiser et réserver une place de parking en ligne')
add_bullet(doc, 'Offrir aux agents un outil de gestion en temps réel des places de stationnement')
add_bullet(doc, 'Fournir aux administrateurs une vue d\'ensemble du système avec des statistiques')
add_bullet(doc, 'Intégrer un système de codes promotionnels pour fidéliser la clientèle')
add_bullet(doc, 'Assurer la communication par email pour les confirmations et notifications')

doc.add_heading('1.3 Périmètre du projet', level=2)
doc.add_paragraph(
    'Le projet couvre le développement d\'une application web complète avec :'
)
add_bullet(doc, 'Un backend API REST développé avec Laravel 12')
add_bullet(doc, 'Un frontend monopage (SPA) développé avec React 19')
add_bullet(doc, 'Une base de données relationnelle MySQL')
add_bullet(doc, 'Une authentification sécurisée par tokens (Sanctum)')
add_bullet(doc, 'Un système d\'envoi d\'emails transactionnels')
add_bullet(doc, 'Une cartographie interactive avec Leaflet')

# ══════════════════════════════════════════════════════════════
# CHAPITRE 2
# ══════════════════════════════════════════════════════════════
doc.add_heading('Chapitre 2 : Analyse et spécification des besoins', level=1)

doc.add_heading('2.1 Besoins fonctionnels', level=2)
doc.add_paragraph('Les besoins fonctionnels identifiés sont organisés par acteur :')

doc.add_heading('Côté Client', level=3)
add_bullet(doc, 'Créer un compte et s\'authentifier')
add_bullet(doc, 'Rechercher des parkings par ville ou par géolocalisation')
add_bullet(doc, 'Consulter la disponibilité des places en temps réel')
add_bullet(doc, 'Sélectionner une place et effectuer une réservation')
add_bullet(doc, 'Gérer ses véhicules (ajout, modification, suppression)')
add_bullet(doc, 'Appliquer des codes promotionnels')
add_bullet(doc, 'Consulter l\'historique des réservations')
add_bullet(doc, 'Annuler une réservation')
add_bullet(doc, 'Modifier son profil')

doc.add_heading('Côté Agent', level=3)
add_bullet(doc, 'Visualiser les statistiques de son parking')
add_bullet(doc, 'Consulter la liste des réservations')
add_bullet(doc, 'Confirmer ou refuser des réservations')
add_bullet(doc, 'Gérer l\'état des places (libre/occupé)')
add_bullet(doc, 'Modifier les informations de son parking')

doc.add_heading('Côté Administrateur', level=3)
add_bullet(doc, 'Consulter les statistiques globales')
add_bullet(doc, 'Gérer les parkings (CRUD)')
add_bullet(doc, 'Gérer les agents (CRUD)')
add_bullet(doc, 'Gérer les codes promotionnels')
add_bullet(doc, 'Envoyer des codes promo par email aux abonnés')
add_bullet(doc, 'Exporter les réservations en CSV')
add_bullet(doc, 'Consulter toutes les réservations')

doc.add_heading('2.2 Besoins non fonctionnels', level=2)
add_bullet(doc, 'Sécurité : Authentification par tokens, protection des routes par middleware')
add_bullet(doc, 'Performance : Temps de réponse rapide pour l\'affichage des places disponibles')
add_bullet(doc, 'Disponibilité : Application accessible 24h/24 et 7j/7')
add_bullet(doc, 'Maintenabilité : Code structuré suivant le pattern MVC')
add_bullet(doc, 'Évolutivité : Architecture modulaire permettant l\'ajout de fonctionnalités')
add_bullet(doc, 'Compatibilité : Application responsive fonctionnant sur tous les navigateurs modernes')
add_bullet(doc, 'Expérience utilisateur : Interface intuitive et moderne avec animations')

doc.add_heading('2.3 Acteurs du système', level=2)
add_table(doc, ['Acteur', 'Rôle', 'Fonctionnalités principales'], [
    ['Client', 'Conducteur recherchant une place', 'Réserver, gérer véhicules, profiter des promos'],
    ['Agent', 'Gestionnaire de parking', 'Gérer les places, confirmer/refuser les réservations'],
    ['Administrateur', 'Superviseur', 'Gérer parkings, agents, promos, statistiques'],
])

doc.add_heading('2.4 Diagramme de cas d\'utilisation', level=2)
doc.add_paragraph('[Diagramme à insérer ici]')

# ══════════════════════════════════════════════════════════════
# CHAPITRE 3
# ══════════════════════════════════════════════════════════════
doc.add_heading('Chapitre 3 : Conception', level=1)

doc.add_heading('3.1 Architecture du système', level=2)
doc.add_paragraph(
    'ParLak suit une architecture REST (API-first) où le frontend React communique avec le backend '
    'Laravel exclusivement via des requêtes HTTP JSON. L\'authentification est gérée par des tokens '
    'Sanctum stockés dans le localStorage du navigateur.'
)
doc.add_paragraph('[Schéma d\'architecture à insérer ici]')

doc.add_heading('Organisation des routes API', level=3)
add_table(doc, ['Catégorie', 'Middleware', 'Nombre de routes'], [
    ['Publiques', 'Aucun', '6'],
    ['Client (auth:sanctum)', 'auth:sanctum', '16'],
    ['Agent (auth:sanctum)', 'auth:sanctum', '6'],
    ['Administrateur', 'auth:sanctum + admin', '16'],
    ['Total', '', '48'],
])

doc.add_heading('Flux de données : réservation', level=3)
doc.add_paragraph(
    '1. Le client sélectionne une ville et un parking dans le dashboard\n'
    '2. Le client choisit une place libre sur le plan interactif\n'
    '3. Le client peut appliquer un code promo (validation via API)\n'
    '4. Confirmation : une réservation est créée avec un QR code unique\n'
    '5. Un email de confirmation est envoyé automatiquement\n'
    '6. L\'agent voit la réservation dans son espace et peut la confirmer/refuser'
)

doc.add_heading('3.2 Modélisation de la base de données', level=2)
doc.add_paragraph(
    'La base de données est composée de 9 tables principales reliées par des clés étrangères.'
)
doc.add_paragraph('[Diagramme de classes UML / MCD à insérer ici]')

doc.add_heading('Tables et colonnes', level=3)

tables_info = [
    ('users', [
        ('id', 'PK'), ('firstname', 'VARCHAR'), ('lastname', 'VARCHAR nullable'),
        ('email', 'VARCHAR UNIQUE'), ('phone', 'VARCHAR nullable'), ('role', 'ENUM client/agent/admin'),
        ('password', 'VARCHAR hashed'),
    ]),
    ('parkings', [
        ('id', 'PK'), ('city_id', 'FK → cities'), ('user_id', 'FK → users nullable'),
        ('name', 'VARCHAR'), ('address', 'VARCHAR'), ('latitude', 'DECIMAL nullable'),
        ('longitude', 'DECIMAL nullable'), ('total_spots', 'INTEGER'), ('is_active', 'BOOLEAN'),
    ]),
    ('parking_zones', [
        ('id', 'PK'), ('parking_id', 'FK → parkings'), ('name', 'VARCHAR'), ('level', 'INTEGER'),
    ]),
    ('parking_spots', [
        ('id', 'PK'), ('zone_id', 'FK → parking_zones'), ('code', 'VARCHAR'),
        ('type', 'ENUM'), ('price_per_hour', 'DECIMAL'), ('status', 'ENUM libre/occupee/reservee'),
    ]),
    ('reservations', [
        ('id', 'PK'), ('reference', 'VARCHAR UNIQUE'), ('user_id', 'FK → users'),
        ('vehicle_id', 'FK → vehicles'), ('spot_id', 'FK → parking_spots'),
        ('entry_date', 'DATE'), ('entry_time', 'TIME'), ('exit_time', 'TIME'),
        ('duration_hours', 'INTEGER'), ('total_price', 'DECIMAL'),
        ('discount_amount', 'DECIMAL nullable'), ('final_price', 'DECIMAL nullable'),
        ('promo_code_id', 'FK nullable'), ('status', 'VARCHAR'), ('qr_token', 'VARCHAR UNIQUE'),
    ]),
    ('vehicles', [
        ('id', 'PK'), ('user_id', 'FK → users'), ('plate', 'VARCHAR'),
        ('make', 'VARCHAR'), ('model', 'VARCHAR'), ('color', 'VARCHAR nullable'), ('type', 'VARCHAR'),
    ]),
    ('promo_codes', [
        ('id', 'PK'), ('code', 'VARCHAR UNIQUE'), ('type', 'ENUM percent/flat'),
        ('discount', 'DECIMAL'), ('max_uses', 'INTEGER nullable'), ('uses_count', 'INTEGER'),
        ('expires_at', 'DATE nullable'), ('is_active', 'BOOLEAN'),
    ]),
    ('cities', [('id', 'PK'), ('name', 'VARCHAR')]),
    ('newsletter_subscribers', [('id', 'PK'), ('email', 'VARCHAR UNIQUE')]),
]

for table_name, columns in tables_info:
    doc.add_heading(f'Table : {table_name}', level=4)
    add_table(doc, ['Colonne', 'Type / Contrainte'], columns)

doc.add_heading('Relations entre les tables', level=3)
add_bullet(doc, 'users 1──N reservations : un utilisateur peut avoir plusieurs réservations')
add_bullet(doc, 'users 1──N vehicles : un utilisateur peut enregistrer plusieurs véhicules')
add_bullet(doc, 'users 1──1 parking : un agent gère un seul parking')
add_bullet(doc, 'parking 1──N zones : un parking contient plusieurs zones')
add_bullet(doc, 'zone 1──N spots : une zone contient plusieurs places')
add_bullet(doc, 'cities 1──N parkings : une ville peut avoir plusieurs parkings')

doc.add_heading('3.3 Diagrammes UML', level=2)
doc.add_paragraph('[Diagramme de séquence — Réservation à insérer ici]')
doc.add_paragraph('[Diagramme de classes — Modèle de données à insérer ici]')
doc.add_paragraph('[Diagramme de déploiement à insérer ici]')

# ══════════════════════════════════════════════════════════════
# CHAPITRE 4
# ══════════════════════════════════════════════════════════════
doc.add_heading('Chapitre 4 : Implémentation', level=1)

doc.add_heading('4.1 Technologies utilisées', level=2)

doc.add_heading('Backend', level=3)
add_table(doc, ['Technologie', 'Version', 'Rôle'], [
    ['PHP', '8.2+', 'Langage serveur'],
    ['Laravel', '12', 'Framework PHP (MVC)'],
    ['Laravel Sanctum', '*', 'Authentification API par tokens'],
    ['MySQL', '8+', 'Base de données relationnelle'],
    ['Gmail SMTP', '-', 'Envoi d\'emails transactionnels'],
])

doc.add_heading('Frontend', level=3)
add_table(doc, ['Technologie', 'Version', 'Rôle'], [
    ['React', '19', 'Bibliothèque UI (SPA)'],
    ['Vite', '7', 'Build tool et serveur de développement'],
    ['Leaflet / react-leaflet', '1.9 / 5.0', 'Cartographie interactive'],
    ['Lucide React', '0.563', 'Iconographie'],
    ['Lottie React', '2.4', 'Animations vectorielles'],
])

doc.add_heading('4.2 Backend (Laravel 12)', level=2)
doc.add_paragraph(
    'Le backend est organisé selon le modèle MVC de Laravel. Il se compose de 8 contrôleurs, '
    '9 modèles, 16 migrations, 1 middleware personnalisé, 4 seeders et 3 classes Mailable.'
)

doc.add_heading('Contrôleurs', level=3)
add_table(doc, ['Contrôleur', 'Méthodes principales', 'Description'], [
    ['AuthController', 'register, login, me, logout', 'Authentification et gestion des sessions'],
    ['VehicleController', 'index, store, update, destroy', 'Gestion des véhicules clients'],
    ['ReservationController', 'index, store, confirm, refuse, cancel', 'Cycle de vie des réservations'],
    ['PromoController', 'index, store, validateCode, apply, sendToSubscribers', 'Codes promotionnels'],
    ['ParkingController', 'index, zones, agentSpots, updateSpot', 'Parkings, zones et places'],
    ['AdminController', 'dashboardStats, parkings, agents, promos', 'Administration générale'],
    ['UserController', 'show, update', 'Gestion du profil utilisateur'],
])

doc.add_paragraph('[Extraits de code et explications détaillées à insérer ici]')

doc.add_heading('Sécurité', level=3)
doc.add_paragraph(
    'L\'authentification est assurée par Laravel Sanctum avec des tokens API. '
    'Le middleware EnsureAdmin vérifie que l\'utilisateur a le rôle "admin" pour les routes '
    'administrateur. Les mots de passe sont hachés avec Bcrypt.'
)

doc.add_heading('Emails transactionnels', level=3)
add_table(doc, ['Classe Mailable', 'Déclencheur', 'Template'], [
    ['WelcomeNewsletter', 'Inscription newsletter', 'emails/welcome_newsletter.blade.php'],
    ['ReservationConfirmed', 'Confirmation de réservation', 'emails/reservation_confirmed.blade.php'],
    ['PromoCodeEmail', 'Envoi promo aux abonnés', 'emails/promo_code.blade.php'],
])

doc.add_heading('4.3 Frontend (React 19)', level=2)
doc.add_paragraph(
    'Le frontend est une application monopage (SPA) développée avec React 19 et Vite 7. '
    'La navigation est gérée par un état React sans routeur externe. L\'application se compose '
    'de 24 composants React, 17 fichiers CSS et 6 animations Lottie.'
)

doc.add_heading('Composants principaux', level=3)
add_table(doc, ['Composant', 'Fonctionnalités', 'Lignes'], [
    ['App.jsx', 'Point d\'entrée, routing, héros, cartographie', '~500'],
    ['ClientDashboard.jsx', 'Réservation, véhicules, profil, paiement', '~1500'],
    ['AgentPage.jsx', 'Places, réservations, paramètres parking', '~900'],
    ['AdminPage.jsx', 'Parkings, agents, promos, statistiques', '~1200'],
    ['Login.jsx', 'Connexion client/agent', '~200'],
    ['AdminLogin.jsx', 'Connexion administrateur', '~100'],
    ['inscription.jsx', 'Inscription client/agent', '~500'],
])

doc.add_paragraph('[Extraits de code et explications détaillées à insérer ici]')

doc.add_heading('Gestion d\'état', level=3)
doc.add_paragraph(
    'Chaque dashboard gère son propre état avec useState et useEffect. '
    'Un helper apiFetch centralise les appels API avec le token d\'authentification.'
)

# ══════════════════════════════════════════════════════════════
# CHAPITRE 5
# ══════════════════════════════════════════════════════════════
doc.add_heading('Chapitre 5 : Fonctionnalités et manuel d\'utilisation', level=1)

doc.add_heading('5.1 Présentation des interfaces', level=2)
doc.add_paragraph('[Captures d\'écran des interfaces à insérer ici]')

doc.add_heading('Interface Client', level=3)
add_bullet(doc, 'Dashboard : tableau de bord avec résumé des activités')
add_bullet(doc, 'Réservation : formulaire de recherche, plan du parking, récapitulatif')
add_bullet(doc, 'Mes Véhicules : liste des véhicules avec CRUD modal')
add_bullet(doc, 'Mes Réservations : historique avec statuts et QR code')
add_bullet(doc, 'Profil : modification des informations personnelles')

doc.add_heading('Interface Agent', level=3)
add_bullet(doc, 'Dashboard : statistiques du parking')
add_bullet(doc, 'Places : grille interactive avec changement de statut')
add_bullet(doc, 'Réservations : liste avec confirmation/refus')
add_bullet(doc, 'Paramètres : modification du parking')

doc.add_heading('Interface Administrateur', level=3)
add_bullet(doc, 'Dashboard : statistiques globales')
add_bullet(doc, 'Parkings : gestion CRUD')
add_bullet(doc, 'Agents : gestion CRUD')
add_bullet(doc, 'Promos : création et envoi aux abonnés')
add_bullet(doc, 'Réservations : vue d\'ensemble')

doc.add_heading('5.2 Manuel d\'utilisation', level=2)

doc.add_heading('Guide Client', level=3)
add_bullet(doc, '1. Créez un compte via le formulaire d\'inscription')
add_bullet(doc, '2. Connectez-vous avec votre email et mot de passe')
add_bullet(doc, '3. Ajoutez un ou plusieurs véhicules dans "Mes Véhicules"')
add_bullet(doc, '4. Allez dans "Réserver une place" et choisissez votre ville')
add_bullet(doc, '5. Sélectionnez un parking, une zone et une place disponible')
add_bullet(doc, '6. Appliquez un code promo si vous en avez un')
add_bullet(doc, '7. Confirmez la réservation et téléchargez votre QR code')
add_bullet(doc, '8. Présentez le QR code à l\'entrée du parking')

doc.add_heading('Guide Agent', level=3)
add_bullet(doc, '1. Créez un compte agent via l\'inscription')
add_bullet(doc, '2. Le tableau de bord affiche les statistiques de votre parking')
add_bullet(doc, '3. Gérez les places (libre/occupée) dans l\'onglet "Places"')
add_bullet(doc, '4. Consultez les réservations et confirmez/refusez')
add_bullet(doc, '5. Modifiez les informations de votre parking dans "Paramètres"')

doc.add_heading('Guide Administrateur', level=3)
add_bullet(doc, '1. Connectez-vous via la page admin dédiée')
add_bullet(doc, '2. Le tableau de bord présente les statistiques globales')
add_bullet(doc, '3. Gérez les parkings (ajout, modification, suppression)')
add_bullet(doc, '4. Gérez les agents et leurs comptes')
add_bullet(doc, '5. Créez et envoyez des codes promotionnels')
add_bullet(doc, '6. Consultez l\'ensemble des réservations')

# ══════════════════════════════════════════════════════════════
# CHAPITRE 6
# ══════════════════════════════════════════════════════════════
doc.add_heading('Chapitre 6 : Tests et validation', level=1)
doc.add_paragraph('[Résultats des tests à insérer ici]')

doc.add_heading('Tests unitaires', level=2)
doc.add_paragraph('Description des tests unitaires effectués sur les contrôleurs et les modèles.')

doc.add_heading('Tests d\'intégration', level=2)
doc.add_paragraph('Tests de bout en bout des principales fonctionnalités.')

doc.add_heading('Validation des fonctionnalités', level=2)
add_table(doc, ['Fonctionnalité', 'Statut', 'Commentaire'], [
    ['Authentification', '✅ Testé', 'Inscription, connexion, déconnexion OK'],
    ['Gestion des véhicules', '✅ Testé', 'CRUD complet OK'],
    ['Réservation', '✅ Testé', 'Création, confirmation, annulation OK'],
    ['Codes promo', '✅ Testé', 'Validation et application OK'],
    ['Newsletter', '✅ Testé', 'Inscription et envoi email OK'],
    ['Envoi d\'emails', '✅ Testé', 'SMTP Gmail fonctionnel'],
    ['Cartographie', '✅ Testé', 'Leaflet et géolocalisation OK'],
])

# ══════════════════════════════════════════════════════════════
# CONCLUSION
# ══════════════════════════════════════════════════════════════
doc.add_heading('Conclusion et perspectives', level=1)

doc.add_paragraph(
    'Le projet ParLak a permis de développer une plateforme complète de gestion de stationnement '
    'intelligent. En combinant Laravel 12 pour le backend et React 19 pour le frontend, nous avons '
    'créé une application moderne, réactive et évolutive. Les trois profils d\'utilisateurs (client, '
    'agent, admin) disposent d\'interfaces adaptées à leurs besoins spécifiques.'
)

doc.add_heading('Compétences acquises', level=2)
add_bullet(doc, 'Maîtrise de Laravel 12 pour la création d\'API REST')
add_bullet(doc, 'Développement d\'applications monopage avec React 19')
add_bullet(doc, 'Gestion de l\'authentification avec Laravel Sanctum')
add_bullet(doc, 'Intégration de cartes interactives avec Leaflet')
add_bullet(doc, 'Conception et optimisation de bases de données relationnelles')
add_bullet(doc, 'Envoi d\'emails transactionnels via SMTP')

doc.add_heading('Difficultés rencontrées', level=2)
add_bullet(doc, 'Configuration SMTP Gmail (mot de passe d\'application)')
add_bullet(doc, 'Gestion des relations complexes entre les tables')
add_bullet(doc, 'Débogage des requêtes asynchrones entre frontend et backend')

doc.add_heading('Perspectives d\'amélioration', level=2)
add_bullet(doc, 'Intégration d\'un système de paiement en ligne (CIB, PayPal)')
add_bullet(doc, 'Application mobile Android/iOS avec React Native')
add_bullet(doc, 'Système de notifications push en temps réel')
add_bullet(doc, 'Programme de fidélité avec points cumulables')
add_bullet(doc, 'Intelligence artificielle pour prédire l\'affluence')
add_bullet(doc, 'Support multilingue avancé (FR/EN/AR)')

# ══════════════════════════════════════════════════════════════
# BIBLIOGRAPHIE
# ══════════════════════════════════════════════════════════════
doc.add_heading('Bibliographie / Webographie', level=1)
add_bullet(doc, 'Documentation Laravel 12 — https://laravel.com/docs')
add_bullet(doc, 'Documentation React 19 — https://react.dev')
add_bullet(doc, 'Documentation Laravel Sanctum — https://laravel.com/docs/sanctum')
add_bullet(doc, 'Leaflet — https://leafletjs.com')
add_bullet(doc, 'React Leaflet — https://react-leaflet.js.org')
add_bullet(doc, 'MySQL — https://dev.mysql.com/doc/')
add_bullet(doc, 'Vite — https://vitejs.dev')
add_bullet(doc, 'Lucide Icons — https://lucide.dev')

# ══════════════════════════════════════════════════════════════
# ANNEXES
# ══════════════════════════════════════════════════════════════
doc.add_heading('Annexes', level=1)
add_para(doc, '[Annexes à insérer ici]', italic=True, size=11)
doc.add_paragraph()
doc.add_paragraph('Annexe 1 : Scripts SQL de création des tables')
doc.add_paragraph('Annexe 2 : Liste complète des routes API')
doc.add_paragraph('Annexe 3 : Code source des contrôleurs principaux')
doc.add_paragraph('Annexe 4 : Configuration de l\'environnement')

doc.add_paragraph()
add_para(doc, '— FIN DU RAPPORT —', bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x04, 0x15, 0x62))

# ── Save ──
output_path = 'C:\\Users\\dell\\Desktop\\MyPfe\\Rapport_PFE_ParLak_v2.docx'
doc.save(output_path)
print(f'Rapport généré avec succès : {output_path}')
